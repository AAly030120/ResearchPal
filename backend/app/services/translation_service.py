import os
import logging
from docx import Document
from app.services.llm_service import llm_service
from app.services.file_parser import extract_text
from app.core.config import settings

logger = logging.getLogger(__name__)


def _chunk_text(text: str, max_chars: int = 3000) -> list[str]:
    """Split text into chunks roughly respecting sentence boundaries."""
    chunks = []
    current = []
    current_len = 0
    for sentence in text.replace('\n', ' ').split('.'):
        s = sentence.strip() + '.'
        if current_len + len(s) > max_chars and current:
            chunks.append(' '.join(current))
            current = []
            current_len = 0
        if s.strip():
            current.append(s)
            current_len += len(s)
    if current:
        chunks.append(' '.join(current))
    return chunks if chunks else [text]


async def translate_text(
    text: str,
    source_lang: str,
    target_lang: str,
    model: str = "gpt-4o-mini",
) -> str:
    """Translate text using LLM."""
    system_prompt = (
        f"You are a professional translator. Translate the following text from "
        f"{source_lang} to {target_lang}. Preserve formatting, tone, and structure. "
        f"Return ONLY the translated text, no explanations."
    )
    messages = [{"role": "user", "content": text}]
    return await llm_service.chat_complete(model, messages, system_prompt, temperature=0.2)


async def translate_file(
    file_id: str,
    file_path: str,
    file_type: str,
    source_lang: str,
    target_lang: str,
    model: str = "gpt-4o-mini",
) -> str:
    """Translate a file and save as docx."""
    parsed = extract_text(file_id, file_path, file_type)
    text = parsed.get("text", "")

    if not text.strip():
        raise ValueError("No text content found in file.")

    chunks = _chunk_text(text, max_chars=3000)
    translated_chunks = []
    for i, chunk in enumerate(chunks):
        logger.info(f"Translating chunk {i+1}/{len(chunks)}")
        translated = await translate_text(chunk, source_lang, target_lang, model)
        translated_chunks.append(translated)

    full_translation = "\n\n".join(translated_chunks)

    # Save as docx
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    import uuid
    filename = f"translated_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.abspath(os.path.join(settings.OUTPUT_DIR, filename))

    doc = Document()
    doc.add_heading(f"Translation: {source_lang} -> {target_lang}", level=1)
    for paragraph in full_translation.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(output_path)
    logger.info(f"Translation saved to {output_path}")
    return output_path
